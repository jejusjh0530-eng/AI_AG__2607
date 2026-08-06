"""General-purpose .docx document generator using python-docx.

Usage:
    python create_document.py [--output PATH] [--title TEXT] [--subtitle TEXT] [--color HEXCODE]

Example:
    python create_document.py --title "Project Report" --color "#2E86AB" --output report.docx

Styling defaults:
    - Font: 맑은 고딕 (Malgun Gothic), applied to Latin and East Asian text alike
    - Title size: 40pt
    - Body size: 20pt
    - Main color: applied to headings and the table header row (customizable via --color)
"""

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_NAME = "맑은 고딕"
TITLE_SIZE = Pt(40)
HEADING_SIZE = Pt(26)
BODY_SIZE = Pt(20)
TABLE_TEXT_SIZE = Pt(16)
DEFAULT_COLOR = "1F4E79"

HEX_PATTERN_LENGTH = 6


def parse_color(value: str) -> RGBColor:
    hex_value = value.lstrip("#")
    if len(hex_value) != HEX_PATTERN_LENGTH or any(
        c not in "0123456789ABCDEFabcdef" for c in hex_value
    ):
        raise argparse.ArgumentTypeError(
            f"'{value}' is not a valid hex color (expected format: RRGGBB or #RRGGBB)"
        )
    return RGBColor.from_string(hex_value.upper())


def set_run_font(run, size: Pt = None, color: RGBColor = None, bold: bool = None):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    return run


def shade_cell(cell, color: RGBColor):
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), str(color))
    cell._tc.get_or_add_tcPr().append(shade)


def add_page_number_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=Pt(10))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


def apply_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_styled_heading(doc: Document, text: str, color: RGBColor):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=HEADING_SIZE, color=color, bold=True)
    return paragraph


def add_styled_paragraph(doc: Document, text: str, style: str = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=BODY_SIZE)
    return paragraph


def build_document(title_text: str, subtitle_text: str, main_color: RGBColor) -> Document:
    doc = Document()

    doc.core_properties.title = title_text
    doc.core_properties.subject = subtitle_text

    apply_default_font(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title_text)
    set_run_font(title_run, size=TITLE_SIZE, color=main_color, bold=True)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle_text)
    set_run_font(subtitle_run, size=BODY_SIZE, color=RGBColor(0x59, 0x59, 0x59))
    subtitle_run.italic = True

    add_styled_heading(doc, "1. Introduction", main_color)
    add_styled_paragraph(
        doc,
        "This is a body paragraph. Replace this text with your own content. "
        "You can add as many paragraphs as needed under each heading.",
    )

    add_styled_heading(doc, "2. Details", main_color)
    for point in ["Key point one.", "Key point two.", "Key point three."]:
        add_styled_paragraph(doc, point, style="List Bullet")

    add_styled_heading(doc, "3. Data", main_color)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Item", "Description", "Value"]
    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=TABLE_TEXT_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        shade_cell(cell, main_color)

    sample_rows = [
        ("Item 1", "Description of item 1", "-"),
        ("Item 2", "Description of item 2", "-"),
        ("Item 3", "Description of item 3", "-"),
    ]
    for item, description, value in sample_rows:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, (item, description, value)):
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, size=TABLE_TEXT_SIZE)

    add_styled_heading(doc, "4. Conclusion", main_color)
    add_styled_paragraph(doc, "Summarize the document here.")

    add_page_number_footer(doc)

    return doc


def parse_args():
    parser = argparse.ArgumentParser(description="Generate a general-purpose .docx document.")
    parser.add_argument("--output", type=Path, default=Path("output.docx"), help="Output file path")
    parser.add_argument("--title", default="Document Title", help="Document title text")
    parser.add_argument(
        "--subtitle",
        default="Subtitle or short description goes here.",
        help="Document subtitle text",
    )
    parser.add_argument(
        "--color",
        type=parse_color,
        default=parse_color(DEFAULT_COLOR),
        help="Main color as hex code, e.g. #2E86AB (default: #1F4E79)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(args.title, args.subtitle, args.color)
    document.save(args.output)
    print(f"Document saved to {args.output.resolve()}")


if __name__ == "__main__":
    main()
