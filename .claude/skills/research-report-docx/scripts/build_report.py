"""Build a styled .docx research report from a structured JSON input.

This generalizes scripts/create_document.py into a data-driven builder:
instead of a fixed template, the document's sections (paragraphs, bullet
lists, tables) are supplied via a JSON file, while the visual styling
(font, sizes, main color) stays consistent and configurable.

Usage:
    python build_report.py --input findings.json --output report.docx [--color "#2E86AB"]

JSON input schema:
{
  "title": "Document title",
  "subtitle": "Optional subtitle",
  "color": "#1F4E79",                 // optional, overridden by --color if given
  "sections": [
    {"heading": "1. Overview", "paragraphs": ["Paragraph text ..."]},
    {"heading": "2. Key Points", "bullets": ["Point one", "Point two"]},
    {"heading": "3. Data", "table": {"headers": ["Item", "Value"],
                                      "rows": [["A", "1"], ["B", "2"]]}},
    {"heading": "4. Sources", "bullets": ["https://example.com/article"]}
  ]
}

Each section may combine "paragraphs", "bullets", and "table" as needed;
any of the three keys may be omitted.

Styling defaults:
    - Font: 맑은 고딕 (Malgun Gothic), applied to Latin and East Asian text alike
    - Title size: 40pt
    - Section heading size: 26pt
    - Body/bullet/table text size: 20pt / 20pt / 16pt
    - Main color: applied to headings and the table header row (customizable)
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_NAME = "맑은 고딕"
TITLE_SIZE = Pt(40)
HEADING_SIZE = Pt(26)
BODY_SIZE = Pt(20)
TABLE_TEXT_SIZE = Pt(16)
DEFAULT_COLOR = "1F4E79"

HEX_LENGTH = 6


def parse_color(value: str) -> RGBColor:
    hex_value = value.lstrip("#")
    if len(hex_value) != HEX_LENGTH or any(
        c not in "0123456789ABCDEFabcdef" for c in hex_value
    ):
        raise ValueError(
            f"'{value}' is not a valid hex color (expected format: RRGGBB or #RRGGBB)"
        )
    return RGBColor.from_string(hex_value.upper())


def set_run_font(run, size: Pt = None, color: RGBColor = None, bold: bool = None):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    return run


def shade_cell(cell, color: RGBColor):
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), str(color))
    cell._tc.get_or_add_tcPr().append(shade)


def add_page_number_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=Pt(10))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


def apply_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_styled_heading(doc: Document, text: str, color: RGBColor):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=HEADING_SIZE, color=color, bold=True)
    return paragraph


def add_styled_paragraph(doc: Document, text: str, style: str = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=BODY_SIZE)
    return paragraph


def add_styled_table(doc: Document, headers, rows, color: RGBColor):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(str(text))
        set_run_font(run, size=TABLE_TEXT_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        shade_cell(cell, color)

    for row in rows:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, row):
            run = cell.paragraphs[0].add_run(str(text))
            set_run_font(run, size=TABLE_TEXT_SIZE)

    return table


def build_report(data: dict, color_override: RGBColor = None) -> Document:
    title_text = data.get("title", "Untitled Report")
    subtitle_text = data.get("subtitle", "")
    main_color = color_override or parse_color(data.get("color", DEFAULT_COLOR))

    doc = Document()
    doc.core_properties.title = title_text
    doc.core_properties.subject = subtitle_text

    apply_default_font(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title_text)
    set_run_font(title_run, size=TITLE_SIZE, color=main_color, bold=True)

    if subtitle_text:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle_text)
        set_run_font(subtitle_run, size=BODY_SIZE, color=RGBColor(0x59, 0x59, 0x59))
        subtitle_run.italic = True

    for section in data.get("sections", []):
        heading = section.get("heading")
        if heading:
            add_styled_heading(doc, heading, main_color)

        for paragraph_text in section.get("paragraphs", []):
            add_styled_paragraph(doc, paragraph_text)

        for bullet_text in section.get("bullets", []):
            add_styled_paragraph(doc, bullet_text, style="List Bullet")

        table_data = section.get("table")
        if table_data:
            add_styled_table(doc, table_data.get("headers", []), table_data.get("rows", []), main_color)

    add_page_number_footer(doc)

    return doc


def parse_args():
    parser = argparse.ArgumentParser(description="Build a styled .docx report from a JSON research summary.")
    parser.add_argument("--input", type=Path, required=True, help="Path to the input JSON file")
    parser.add_argument("--output", type=Path, default=Path("report.docx"), help="Output .docx path")
    parser.add_argument(
        "--color",
        type=parse_color,
        default=None,
        help="Main color as hex code, e.g. #2E86AB (overrides the JSON's 'color' field)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = build_report(data, color_override=args.color)
    document.save(args.output)
    print(f"Document saved to {args.output.resolve()}")


if __name__ == "__main__":
    main()
