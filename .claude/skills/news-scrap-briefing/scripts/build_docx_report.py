"""Build a styled .docx briefing report from a structured JSON input.

Same data-driven shape as the sibling research-report-docx skill's
build_report.py, extended with two things this skill needs:

  - an "image" key per section, for embedding an article screenshot
    right under its heading/summary
  - an "include_in" key per section, so a JSON built once can drive
    both this docx report and the sibling pptx deck, while letting
    some sections (e.g. a long "Sources" list) appear in only one of
    the two outputs

Usage:
    python build_docx_report.py --input briefing.json --output briefing.docx [--color "#2E86AB"]

JSON input schema:
{
  "title": "Document title",
  "subtitle": "Optional subtitle (e.g. date or scope)",
  "color": "#1F4E79",                 // optional, overridden by --color if given
  "sections": [
    {"heading": "1. Overview", "paragraphs": ["..."]},
    {"heading": "Article headline", "paragraphs": ["summary ..."],
     "bullets": ["key point 1", "key point 2"],
     "image": "path/to/screenshot.png", "source": "Outlet · https://..."},
    {"heading": "Sources", "bullets": ["https://..."], "include_in": ["docx"]}
  ]
}

Each section may combine "paragraphs", "bullets", "image", and "table" as
needed; any key may be omitted. A section is skipped only if it has an
"include_in" list that does NOT contain "docx" — omit "include_in"
entirely to include a section in both outputs.
"""

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from image_utils import fit_within_box, prepare_image_for_embed, resolve_image_path

FONT_NAME = "맑은 고딕"
TITLE_SIZE = Pt(40)
HEADING_SIZE = Pt(26)
BODY_SIZE = Pt(20)
CAPTION_SIZE = Pt(14)
TABLE_TEXT_SIZE = Pt(16)
DEFAULT_COLOR = "1F4E79"
IMAGE_MAX_WIDTH_IN = 5.5
IMAGE_MAX_HEIGHT_IN = 7.5

HEX_LENGTH = 6


def parse_color(value: str) -> RGBColor:
    hex_value = value.lstrip("#")
    if len(hex_value) != HEX_LENGTH or any(
        c not in "0123456789ABCDEFabcdef" for c in hex_value
    ):
        raise ValueError(
            f"'{value}' is not a valid hex color (expected format: RRGGBB or #RRGGBB)"
        )
    return RGBColor.from_string(hex_value.upper())


def set_run_font(run, size: Pt = None, color: RGBColor = None, bold: bool = None, italic: bool = None):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    return run


def shade_cell(cell, color: RGBColor):
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), str(color))
    cell._tc.get_or_add_tcPr().append(shade)


def add_page_number_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=Pt(10))

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


def apply_default_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_styled_heading(doc: Document, text: str, color: RGBColor):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=HEADING_SIZE, color=color, bold=True)
    return paragraph


def add_styled_paragraph(doc: Document, text: str, style: str = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=BODY_SIZE)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=CAPTION_SIZE, color=RGBColor(0x59, 0x59, 0x59), italic=True)
    return paragraph


def add_styled_table(doc: Document, headers, rows, color: RGBColor):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for cell, text in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(str(text))
        set_run_font(run, size=TABLE_TEXT_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        shade_cell(cell, color)

    for row in rows:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, row):
            run = cell.paragraphs[0].add_run(str(text))
            set_run_font(run, size=TABLE_TEXT_SIZE)

    return table


def add_image(doc: Document, image_path: str, base_dir: Path):
    resolved = resolve_image_path(image_path, base_dir)
    if not resolved.exists():
        add_caption(doc, f"[Screenshot not found: {image_path}]")
        return
    embed_path, width_px, height_px = prepare_image_for_embed(resolved, base_dir)
    width_in, height_in = fit_within_box(width_px, height_px, IMAGE_MAX_WIDTH_IN, IMAGE_MAX_HEIGHT_IN)
    doc.add_picture(str(embed_path), width=Inches(width_in), height=Inches(height_in))
    picture_paragraph = doc.paragraphs[-1]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def section_applies(section: dict, output: str) -> bool:
    include_in = section.get("include_in")
    return include_in is None or output in include_in


def build_report(data: dict, base_dir: Path, color_override: RGBColor = None) -> Document:
    title_text = data.get("title", "Untitled Briefing")
    subtitle_text = data.get("subtitle", "")
    main_color = color_override or parse_color(data.get("color", DEFAULT_COLOR))

    doc = Document()
    doc.core_properties.title = title_text
    doc.core_properties.subject = subtitle_text

    apply_default_font(doc)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title_text)
    set_run_font(title_run, size=TITLE_SIZE, color=main_color, bold=True)

    if subtitle_text:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_paragraph.add_run(subtitle_text)
        set_run_font(subtitle_run, size=BODY_SIZE, color=RGBColor(0x59, 0x59, 0x59))
        subtitle_run.italic = True

    for section in data.get("sections", []):
        if not section_applies(section, "docx"):
            continue

        heading = section.get("heading")
        if heading:
            add_styled_heading(doc, heading, main_color)

        source = section.get("source")
        if source:
            add_caption(doc, source)

        for paragraph_text in section.get("paragraphs", []):
            add_styled_paragraph(doc, paragraph_text)

        for bullet_text in section.get("bullets", []):
            add_styled_paragraph(doc, bullet_text, style="List Bullet")

        image_path = section.get("image")
        if image_path:
            add_image(doc, image_path, base_dir)

        table_data = section.get("table")
        if table_data:
            add_styled_table(doc, table_data.get("headers", []), table_data.get("rows", []), main_color)

    add_page_number_footer(doc)

    return doc


def parse_args():
    parser = argparse.ArgumentParser(description="Build a styled .docx briefing report from a JSON summary.")
    parser.add_argument("--input", type=Path, required=True, help="Path to the input JSON file")
    parser.add_argument("--output", type=Path, default=Path("briefing.docx"), help="Output .docx path")
    parser.add_argument(
        "--color",
        type=parse_color,
        default=None,
        help="Main color as hex code, e.g. #2E86AB (overrides the JSON's 'color' field)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8"))
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = build_report(data, base_dir=args.input.parent, color_override=args.color)
    document.save(args.output)
    print(f"Document saved to {args.output.resolve()}")


if __name__ == "__main__":
    main()
